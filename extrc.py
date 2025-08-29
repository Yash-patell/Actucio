import fitz  
import pdfplumber
from paddleocr import PaddleOCR
from PIL import Image
import io
import requests
import json
import os
from docx import Document          # For Word template filling
from docx2pdf import convert       # Optional: Word → PDF conversion
import zipfile   
from dotenv import load_dotenv
import os
def extract_text(file_path: str):
    ext = os.path.splitext(file_path)[1].lower()
    if ext in [".jpg", ".jpeg", ".png"]:
        return extract_text_from_image(file_path)
    elif ext == ".pdf":
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# Initialize OCR (French + English)
# ocr = PaddleOCR(lang='fr')  # 'fr' = French, add 'en' if bilingual
ocr = PaddleOCR(use_angle_cls=True, lang='fr')


def extract_text_from_image(img_path):
    """./Capture d’écran 2025-07-27 163230 (1).png"""
    results = ocr.ocr(img_path, cls=True)
    text = "\n".join([line[1][0] for line in results[0]]) if results else ""
    return text

def extract_text_from_pdf(pdf_path):
    """Try text extraction first; fallback to OCR if scanned PDF"""
    text_content = ""

    # --- First try native text extraction ---
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            extracted = page.extract_text()
            if extracted:
                text_content += extracted + "\n"

    # --- If no text found, fallback to OCR ---
    import numpy as np
    if not text_content.strip():
        print("⚠️ PDF has no extractable text, running OCR...")
        doc = fitz.open(pdf_path)
        for page_num in range(len(doc)):
            pix = doc[page_num].get_pixmap(dpi=300)
            img_bytes = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
            img_np = np.array(img)

            results = ocr.ocr(img_np, cls=True)

            page_text = ""
            # Flatten OCR results safely
            for res in results:
                for line in res:
                    page_text += line[1][0] + "\n"
            text_content += page_text

    return text_content



# --- Gemini API Helper Functions ---

def build_gemini_prompt(extracted_text):
    """Builds a simple prompt for the LLM to extract data."""
    prompt = f"""
You are a professional document parser.
Extract the fields defined in the provided schema from the text below.
If a field is missing, fill it with "N/A".

Text:
{extracted_text}
"""
    return prompt


def call_gemini_api(prompt, api_key):
    """
    Calls the Gemini API with a strict JSON schema to ensure a reliable output format.
    """
    url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key=" + api_key

    # Define the JSON schema for the desired output
    response_schema = {
        "type": "OBJECT",
        "properties": {
            "denomination_sociale_de_la_sas": {"type": "STRING"},
            "montant_du_capital_social": {"type": "STRING"},
            "adresse_du_siege_social": {"type": "STRING"},
            "ville_du_greffe_rcs": {"type": "STRING"},
            "numero_immatriculation_rcs": {"type": "STRING"},
            "date_cloture_exercice": {"type": "STRING"},
            "contenu_rapport_gestion": {"type": "STRING"},
            "contenu_rapport_conventions_reglementees": {"type": "STRING"},
            "montant_resultat_benefice_ou_perte": {"type": "STRING"},
            "decision_affectation_resultat": {"type": "STRING"},
            "date_assemblee_generale_ag": {"type": "STRING"},
            "heure_assemblee_generale": {"type": "STRING"},
            "lieu_assemblee_generale": {"type": "STRING"},
            "date_convocation_ag": {"type": "STRING"},
            "forme_de_la_convocation": {"type": "STRING"},
            "nom_prenom_president_de_seance": {"type": "STRING"},
            "nom_prenom_secretaire_de_seance": {"type": "STRING"},
            "informations_quorum_representation": {"type": "STRING"},
            "numeros_articles_statuts_pertinents": {"type": "STRING"},
            "nom_prenom_qualite_president_representant_legal": {"type": "STRING"},
            "nom_prenom_denomination_adresse_associe": {"type": "STRING"},
            "nom_prenom_denomination_adresse_cac": {"type": "STRING"},
            "nom_prenom_adresse_delegues_cse": {"type": "STRING"},
            "identite_complete_mandant_procuration": {"type": "STRING"},
            "identite_complete_mandataire_procuration": {"type": "STRING"},
            "lieu_date_redaction_signature_documents": {"type": "STRING"},
            "adresse_greffe": {"type": "STRING"},
            "montant_frais_depot_greffe": {"type": "STRING"}
        }
    }

    payload = {
        "contents": [
            {
                "parts": [{"text": prompt}]
            }
        ],
        "generationConfig": {
            "maxOutputTokens": 2000,
            "responseMimeType": "application/json",
            "responseSchema": response_schema
        }
    }

    headers = {
        "Content-Type": "application/json"
    }

    try:
        response = requests.post(url, headers=headers, json=payload, timeout=60) # Increased timeout
        response.raise_for_status() 
        result = response.json()

        # Check for the inconsistent nested 'parts' key
        content = result['candidates'][0]['content']
        if 'parts' in content:
            # If the model returned text with a JSON string inside, parse it
            json_string = content['parts'][0]['text']
            # Adding a basic retry loop for JSON parsing errors
            for _ in range(3): # Try up to 3 times
                try:
                    extracted_data = json.loads(json_string)
                    return extracted_data
                except json.JSONDecodeError as e:
                    print(f"⚠️ Retrying JSON parsing after error: {e}")
                    # Simple fix for common unterminated string errors
                    json_string = json_string.strip().rstrip('},] ').rstrip(',') + '}'
                    if not json_string.endswith('}'):
                        json_string += '}'
                    continue
            return {} # Return an empty dictionary if parsing fails after retries
        else:
            # Otherwise, the content is the JSON object directly
            extracted_data = content
            
        return extracted_data
    except requests.exceptions.RequestException as e:
        print(f"⚠️ Gemini API Request error: {e}")
        return None
    except (KeyError, IndexError) as e:
        print(f"⚠️ Failed to parse Gemini output: {e}")
        return None
        
def fill_word_template(template_path, output_path, data):
    """
    Fills a Word template with data from the JSON output.
    Placeholders in the template should be in the format: {variable_name}
    """
    try:
        document = Document(template_path)
        # Find and replace in paragraphs
        for p in document.paragraphs:
            for k, v in data.items():
                p.text = p.text.replace('{' + k + '}', str(v))
        # Find and replace in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for k, v in data.items():
                        if '{' + k + '}' in cell.text:
                            cell.text = cell.text.replace('{' + k + '}', str(v))
        
        document.save(output_path)
        return output_path
    except Exception as e:
        print(f"⚠️ Error filling Word template {template_path}: {e}")
        return None
        
def generate_documents(structured_data_all_files, template_path, output_folder):
    """
    Generates Word and PDF documents from the structured data using a single template.
    """
    if not os.path.exists(template_path):
        print(f"⚠️ Template file not found: {template_path}")
        return

    for file_name, data in structured_data_all_files.items():
        print(f"\n📄 Generating documents for {file_name} using template {os.path.basename(template_path)}...")
        
        # Create output paths
        base_name = os.path.splitext(file_name)[0]
        filled_docx_path = os.path.join(output_folder, f"{base_name}_filled.docx")
        pdf_path = os.path.join(output_folder, f"{base_name}_filled.pdf")

        filled_docx_path = fill_word_template(template_path, filled_docx_path, data)
        if filled_docx_path:
            try:
                convert(filled_docx_path, pdf_path)
                print(f"✅ Documents saved: {filled_docx_path} and {pdf_path}")
            except Exception as e:
                print(f"⚠️ Error converting to PDF for {filled_docx_path}: {e}")

if __name__ == "__main__":
    upload_folder = "uploaded_files"
    template_folder = "templates"
    output_folder = "extracted_results"
    
    os.makedirs(upload_folder, exist_ok=True)
    os.makedirs(template_folder, exist_ok=True)
    os.makedirs(output_folder, exist_ok=True)

    uploaded_files = [
        os.path.join(upload_folder, f) 
        for f in os.listdir(upload_folder) 
        if f.lower().endswith((".pdf", ".png", ".jpg", ".jpeg"))
    ]
    if not uploaded_files:
        print("⚠️ No files found in the upload folder!")
        exit()

    all_texts = {}
    for file_path in uploaded_files:
        file_name = os.path.basename(file_path)
        print(f"\n=== Extracting from {file_name} ===")
        text = extract_text(file_path)
        all_texts[file_name] = text
        txt_path = os.path.join(output_folder, f"{os.path.splitext(file_name)[0]}_extracted.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"✅ Saved extracted text to {txt_path}")
    print("\n✅ All extractions done. Ready for structured LLM processing.")

    load_dotenv() 
    api_key = os.getenv("APIkey_gemini")
    structured_data_all_files = {}

    for file_name, text in all_texts.items():
        print(f"\n=== Sending {file_name} to Gemini ===")
        prompt = build_gemini_prompt(text)
        extracted_data = call_gemini_api(prompt, api_key)
        if extracted_data:
            structured_data_all_files[file_name] = extracted_data
            print(f"✅ Extracted structured data for {file_name}")
        else:
            print(f"⚠️ Failed to extract data for {file_name}")
    print("\n\n=== FINAL STRUCTURED OUTPUTS ===")
    for file_name, data in structured_data_all_files.items():
        print(f"\n--- Output for {file_name} ---")
        print(json.dumps(data, indent=2, ensure_ascii=False))
    
    output_json_file = os.path.join(output_folder, "structured_output.json")
    with open(output_json_file, "w", encoding="utf-8") as f:
        json.dump(structured_data_all_files, f, indent=2, ensure_ascii=False)
    print(f"\n✅ All structured data saved to {output_json_file}")
    
    template_name = "template.docx" # Assuming this is your single template file
    template_path = os.path.join(template_folder, template_name)
    generate_documents(structured_data_all_files, template_path, output_folder)
